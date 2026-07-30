import io

from docx import Document

from app.models.candidate import Candidate
from app.models.job import Job, JobStatus
from app.services import claude_service


def _valid_resume_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Jane Doe - Senior Backend Engineer")
    doc.add_paragraph("5 years experience with Python, FastAPI, PostgreSQL, and distributed systems.")
    doc.add_paragraph("Led the payments infrastructure team at Acme Corp.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_open_job(db_session) -> Job:
    job = Job(
        owner_org_id="org_test123",
        title="Senior Backend Engineer",
        description="Build scalable APIs",
        required_skills="Python, FastAPI",
        min_years_experience=3,
        status=JobStatus.OPEN,
    )
    db_session.add(job)
    db_session.commit()
    db_session.refresh(job)
    return job


class TestDeleteCandidate:
    def test_delete_candidate_success(self, client, db_session, mocker):
        job = _make_open_job(db_session)
        mocker.patch("app.services.storage_service.upload_resume", return_value="path.docx")
        mocker.patch("app.services.claude_service.score_resume", side_effect=Exception("skip scoring"))

        apply_resp = client.post(
            "/api/v1/candidates/apply",
            data={"job_id": job.id, "full_name": "Delete Me", "email": "deleteme@example.com"},
            files={"resume": ("resume.docx", _valid_resume_bytes(), "application/octet-stream")},
        )
        candidate_id = apply_resp.json()["id"]

        from app.core.auth import get_current_user, AuthenticatedUser
        from app.main import app as fastapi_app

        fastapi_app.dependency_overrides[get_current_user] = lambda: AuthenticatedUser(
            user_id="u1", org_id="org_test123", claims={}
        )
        try:
            response = client.delete(f"/api/v1/candidates/{candidate_id}")
            assert response.status_code == 204

            # Confirm it's actually gone, not just hidden
            get_response = client.get(f"/api/v1/candidates/{candidate_id}/resume-score")
            assert get_response.status_code == 404
        finally:
            fastapi_app.dependency_overrides.pop(get_current_user, None)

    def test_delete_candidate_org_isolation(self, client, db_session):
        job = _make_open_job(db_session)
        candidate = Candidate(
            job_id=job.id, full_name="Other Org Candidate", email="x@example.com", resume_storage_path="p.pdf"
        )
        db_session.add(candidate)
        db_session.commit()

        from app.core.auth import get_current_user, AuthenticatedUser
        from app.main import app as fastapi_app

        fastapi_app.dependency_overrides[get_current_user] = lambda: AuthenticatedUser(
            user_id="u1", org_id="org_DIFFERENT", claims={}
        )
        try:
            response = client.delete(f"/api/v1/candidates/{candidate.id}")
            assert response.status_code == 404
        finally:
            fastapi_app.dependency_overrides.pop(get_current_user, None)


class TestAuditLog:
    def test_resume_scoring_creates_audit_entry(self, client, db_session, mocker):
        job = _make_open_job(db_session)
        mocker.patch("app.services.storage_service.upload_resume", return_value="path.docx")
        mocker.patch("app.services.email_service.send_interview_invite", return_value=True)
        mocker.patch(
            "app.services.claude_service.score_resume",
            return_value=claude_service.ResumeScoreResult(
                tech_score=88, communication_score=75, role_match_score=90, confidence=85,
                summary="Strong candidate.", strengths="Python depth.", concerns="",
                raw_response="{}", model_version="claude-sonnet-4-6",
            ),
        )
        apply_resp = client.post(
            "/api/v1/candidates/apply",
            data={"job_id": job.id, "full_name": "Jane Doe", "email": "jane@example.com"},
            files={"resume": ("resume.docx", _valid_resume_bytes(), "application/octet-stream")},
        )
        candidate_id = apply_resp.json()["id"]

        from app.core.auth import get_current_user, AuthenticatedUser
        from app.main import app as fastapi_app

        fastapi_app.dependency_overrides[get_current_user] = lambda: AuthenticatedUser(
            user_id="u1", org_id="org_test123", claims={}
        )
        try:
            response = client.get(f"/api/v1/candidates/{candidate_id}/audit-log")
            assert response.status_code == 200
            entries = response.json()
            # Score of 90 is above the default AUTO_INVITE_SCORE_THRESHOLD
            # (60), so this candidate should have BOTH the resume-scoring
            # entry AND an auto-invite entry — this is the audit trail
            # for an automated decision, not a bug.
            assert len(entries) == 2
            step_names = {e["step_name"] for e in entries}
            assert step_names == {"resume_scoring", "auto_invite_threshold"}

            scoring_entry = next(e for e in entries if e["step_name"] == "resume_scoring")
            assert scoring_entry["step_type"] == "probabilistic"
            assert scoring_entry["confidence"] == 85

            invite_entry = next(e for e in entries if e["step_name"] == "auto_invite_threshold")
            assert invite_entry["step_type"] == "deterministic"
            assert invite_entry["confidence"] is None
        finally:
            fastapi_app.dependency_overrides.pop(get_current_user, None)

    def test_low_score_does_not_trigger_auto_invite(self, client, db_session, mocker):
        job = _make_open_job(db_session)
        mocker.patch("app.services.storage_service.upload_resume", return_value="path.docx")
        invite_mock = mocker.patch("app.services.email_service.send_interview_invite", return_value=True)
        mocker.patch(
            "app.services.claude_service.score_resume",
            return_value=claude_service.ResumeScoreResult(
                tech_score=30, communication_score=40, role_match_score=35, confidence=80,
                summary="Weak match.", strengths="", concerns="Limited relevant experience.",
                raw_response="{}", model_version="claude-sonnet-4-6",
            ),
        )
        apply_resp = client.post(
            "/api/v1/candidates/apply",
            data={"job_id": job.id, "full_name": "Weak Candidate", "email": "weak@example.com"},
            files={"resume": ("resume.docx", _valid_resume_bytes(), "application/octet-stream")},
        )
        assert apply_resp.json()["stage"] == "screened"  # not "interview_scheduled"
        invite_mock.assert_not_called()


class TestApplyToJob:
    def test_rejects_application_to_nonexistent_job(self, client):
        response = client.post(
            "/api/v1/candidates/apply",
            data={"job_id": "00000000-0000-0000-0000-000000000000", "full_name": "Jane Doe", "email": "jane@example.com"},
            files={"resume": ("resume.docx", _valid_resume_bytes(), "application/octet-stream")},
        )
        assert response.status_code == 404
        assert response.json()["error_code"] == "not_found"

    def test_rejects_application_to_closed_job(self, client, db_session):
        job = Job(
            owner_org_id="org_test123",
            title="Closed Role",
            description="desc",
            required_skills="Python",
            min_years_experience=1,
            status=JobStatus.CLOSED,
        )
        db_session.add(job)
        db_session.commit()

        response = client.post(
            "/api/v1/candidates/apply",
            data={"job_id": job.id, "full_name": "Jane Doe", "email": "jane@example.com"},
            files={"resume": ("resume.docx", _valid_resume_bytes(), "application/octet-stream")},
        )
        assert response.status_code == 422
        assert "no longer accepting" in response.json()["message"]

    def test_rejects_invalid_email(self, client, db_session):
        job = _make_open_job(db_session)
        response = client.post(
            "/api/v1/candidates/apply",
            data={"job_id": job.id, "full_name": "Jane Doe", "email": "not-an-email"},
            files={"resume": ("resume.docx", _valid_resume_bytes(), "application/octet-stream")},
        )
        assert response.status_code == 422

    def test_successful_application_with_scoring(self, client, db_session, mocker):
        job = _make_open_job(db_session)

        mocker.patch(
            "app.services.storage_service.upload_resume",
            return_value="fake/storage/path.docx",
        )
        mocker.patch(
            "app.services.claude_service.score_resume",
            return_value=claude_service.ResumeScoreResult(
                tech_score=88,
                communication_score=75,
                role_match_score=90,
                summary="Strong candidate.",
                strengths="Python depth.",
                concerns="",
                raw_response="{}",
                model_version="claude-sonnet-4-6",
            ),
        )

        response = client.post(
            "/api/v1/candidates/apply",
            data={"job_id": job.id, "full_name": "Jane Doe", "email": "jane@example.com"},
            files={"resume": ("resume.docx", _valid_resume_bytes(), "application/octet-stream")},
        )
        assert response.status_code == 201
        body = response.json()
        assert body["full_name"] == "Jane Doe"
        assert body["stage"] == "screened"

    def test_application_survives_scoring_failure(self, client, db_session, mocker):
        """
        The candidate's application must be saved even if Claude scoring
        fails downstream — see the comment in candidates.py for why we
        never penalize the candidate for an AI outage.
        """
        job = _make_open_job(db_session)
        mocker.patch(
            "app.services.storage_service.upload_resume",
            return_value="fake/storage/path.docx",
        )
        mocker.patch(
            "app.services.claude_service.score_resume",
            side_effect=Exception("Claude is down"),
        )

        response = client.post(
            "/api/v1/candidates/apply",
            data={"job_id": job.id, "full_name": "Jane Doe", "email": "jane@example.com"},
            files={"resume": ("resume.docx", _valid_resume_bytes(), "application/octet-stream")},
        )
        assert response.status_code == 201
        assert response.json()["stage"] == "applied"
